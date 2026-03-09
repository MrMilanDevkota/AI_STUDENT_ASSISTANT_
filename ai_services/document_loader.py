"""
Document Loader Service
Extracts raw text from uploaded study materials (PDF, DOCX, TXT).
Supports: pypdf for PDFs, python-docx for Word files, plain read for TXT.
"""

import os
import logging

logger = logging.getLogger(__name__)


def load_pdf(file_path: str) -> str:
    """
    Extract text from a PDF file using pypdf.
    Iterates through all pages and concatenates text.
    """
    try:
        from pypdf import PdfReader

        reader = PdfReader(file_path)
        text_parts = []

        for page_num, page in enumerate(reader.pages):
            try:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
            except Exception as e:
                logger.warning(f"Could not extract text from page {page_num}: {e}")
                continue

        extracted = '\n\n'.join(text_parts)

        if not extracted.strip():
            raise ValueError("No text could be extracted from this PDF. It may be scanned/image-based.")

        return extracted

    except ImportError:
        raise ImportError("pypdf is not installed. Run: pip install pypdf")
    except Exception as e:
        logger.error(f"PDF loading error for {file_path}: {e}")
        raise


def load_docx(file_path: str) -> str:
    """
    Extract text from a Word (.docx) file using python-docx.
    Extracts text from paragraphs and table cells.
    """
    try:
        from docx import Document

        doc = Document(file_path)
        text_parts = []

        # Extract from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)

        extracted = '\n\n'.join(text_parts)

        if not extracted.strip():
            raise ValueError("No text found in this Word document.")

        return extracted

    except ImportError:
        raise ImportError("python-docx is not installed. Run: pip install python-docx")
    except Exception as e:
        logger.error(f"DOCX loading error for {file_path}: {e}")
        raise


def load_txt(file_path: str) -> str:
    """
    Read text from a plain text (.txt) file.
    Tries UTF-8 first, falls back to latin-1 encoding.
    """
    try:
        # Try UTF-8 first
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()
    except UnicodeDecodeError:
        # Fallback to latin-1 for files with special characters
        with open(file_path, 'r', encoding='latin-1') as f:
            text = f.read()
    except Exception as e:
        logger.error(f"TXT loading error for {file_path}: {e}")
        raise

    if not text.strip():
        raise ValueError("This text file is empty.")

    return text


def load_document(file_path: str) -> str:
    """
    Main entry point: detect file type and call appropriate loader.

    Args:
        file_path: Absolute path to the uploaded file.

    Returns:
        Extracted text as a single string.

    Raises:
        ValueError: If file type is unsupported or text extraction fails.
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")

    _, extension = os.path.splitext(file_path.lower())

    loaders = {
        '.pdf': load_pdf,
        '.docx': load_docx,
        '.txt': load_txt,
    }

    if extension not in loaders:
        raise ValueError(f"Unsupported file type: {extension}. Allowed: PDF, DOCX, TXT")

    logger.info(f"Loading document: {file_path} (type: {extension})")
    text = loaders[extension](file_path)

    logger.info(f"Successfully extracted {len(text)} characters from {os.path.basename(file_path)}")
    return text
